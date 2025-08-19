import requests
import json
from flask import Flask, request, jsonify
from flask_cors import CORS  # Importing flask_cors to handle CORS issues
from googletrans import Translator
import asyncio
from database import db, migrate
from models import User, Document
from dotenv import load_dotenv
import os
from flask_bcrypt import Bcrypt
from flask_jwt_extended import JWTManager, create_access_token
import re
from datetime import timedelta
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from flask_mail import Mail, Message
from flask_jwt_extended import decode_token
import PyPDF2
import io
import base64
import re
from datetime import datetime
from docx import Document as DocxDocument
import hashlib
import logging
from gemini_setup import get_gemini_response
from gemini_setup import KNOWLEDGE_BASE
import traceback
from fuzzywuzzy import process


load_dotenv()


translator = Translator()

# Initialize logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    handlers=[logging.StreamHandler()]
)

# Initialize the Flassk Application
app = Flask(__name__)

# Enable CORS for all domains (allow React to make requests to this backend)
CORS(app)



# Configure database (SQLite for now)
app.config["SQLALCHEMY_DATABASE_URI"] =  os.getenv("DATABASE_URL")
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False

# Config
app.config["SECRET_KEY"] = os.getenv("SECRET_KEY")
app.config["JWT_SECRET_KEY"] = os.getenv("JWT_SECRET_KEY")

# Configure JWT for authentication
bcrypt = Bcrypt(app) 
jwt = JWTManager(app)
mail = Mail(app)

# Initialize database with app
db.init_app(app)
migrate.init_app(app, db)

# WhatsApp Cloud API credentials
WHATSAPP_PHONE_NUMBER_ID = os.getenv("WHATSAPP_PHONE_NUMBER_ID")
WHATSAPP_ACCESS_TOKEN = os.getenv("WHATSAPP_ACCESS_TOKEN")
VERIFY_TOKEN = os.getenv("VERIFY_TOKEN")

# Configure Flask-Mail for sending emails
app.config["MAIL_SERVER"] = os.getenv("MAIL_SERVER")
app.config["MAIL_PORT"] = int(os.getenv("MAIL_PORT", 587))
app.config["MAIL_USE_TLS"] = os.getenv("MAIL_USE_TLS", "True") == "True"
app.config["MAIL_USERNAME"] = os.getenv("MAIL_USERNAME")
app.config["MAIL_PASSWORD"] = os.getenv("MAIL_PASSWORD")

# Rate limiter (based on client IP)
limiter = Limiter(
    key_func=get_remote_address,
    app=app,
    default_limits=["200 per day", "50 per hour"]  # global fallback
)

def search_knowledge_base(user_message: str):
    """Search across all KB languages using fuzzy matching."""
    best_match = None
    best_score = 0
    best_answer = None

    for lang, qa_pairs in KNOWLEDGE_BASE.items():
        questions = [qa['q'] for qa in qa_pairs]
        match, score = process.extractOne(user_message, questions)

        if score > best_score:
            best_score = score
            best_match = match
            best_answer = next(item['a'] for item in qa_pairs if item['q'] == match)

    logging.info(f"Best KB Match: {best_match} (Score: {best_score})")

    if best_score >= 100:  
        return best_answer
    return None

# --- Validation Helpers ---
def is_valid_email(email):
    return re.match(r"^[^@]+@[^@]+\.[^@]+$", email)

def is_strong_password(password):
    return (
        len(password) >= 8
        and any(c.isupper() for c in password)
        and any(c.islower() for c in password)
        and any(c.isdigit() for c in password)
        and any(c in "!@#$%^&*()-_+=" for c in password)
    )

def is_valid_name(name):
    # Allow only alphabets (first and last names), no emojis, no digits
    return re.match(r"^[A-Za-z]+$", name)


# Pre-set bot replies in multiple languages
BOT_RESPONSES = {
    "verify documents": {
        "en": "Property Verification! Upload your documents for instant analysis.",
        "pidgin": "Bring di paper make we check am sharp sharp for you.",
        "igbo": "Nyochaa ihe osise ụlọ gị! Gị nwere ike nyochaa akwụkwọ gị ozugbo.",
        "yoruba": "Iṣeduro Iwe-ọwọ! Gbe awọn iwe rẹ silẹ fun itupalẹ lẹsẹkẹsẹ.",
        "hausa": "Binciken Dukiya! Aika da takardunku domin duba nan take."
    },
    "report scam": {
        "en": "To report a scam, please provide details via WhatsApp at +234 906 576 0546 or contact support for business.",
        "pidgin": "If person don scam you, yarn us di tori for WhatsApp +234 906 576 0546 or hala support.",
        "igbo": "Iwepụta egwu, biko nye anyị nkọwa site na WhatsApp +234 906 576 0546 ma ọ bụ kpọtụrụ nkwado maka azụmahịa.",
        "yoruba": "Lati ṣe iroyin itanjẹ, jọwọ pese awọn alaye nipasẹ WhatsApp ni +234 906 576 0546 tabi kan si atilẹyin fun iṣowo.",
        "hausa": "Don bayar da rahoton zamba, don Allah samar da cikakkun bayanai ta WhatsApp a +234 906 576 0546 ko tuntuɓi goyon baya don kasuwanci."
    },
    "know your rights": {
        "en": "I can help with Nigeria property laws. Ask me about buyer/renter rights, and I'll guide you!",
        "pidgin": "I sabi Nigerian property law well well. Ask me about buyer or tenant rights make I show you road.",
        "igbo": "Enyemaka na iwu ụlọ Nigeria. Jụọ m gbasara ikike ndị na-azụ ahịa/ndị na-azụ ụlọ, m ga-eduga gị!",
        "yoruba": "Mo le ran ọ lọwọ pẹlu ofin ohun-ini Nigeria. Beere lọwọ mi nipa awọn ẹtọ olura/tabi onílé, emi yoo tọ ọ!",
        "hausa": "Zan iya taimaka muku da dokokin kadarori na Najeriya. Tambayi ni game da hakkin mai saye/matar haya, zan jagorance ku!"
    },
    "contact support": {
        "en": "Contact our Support team at support@trustnet.com or call +234 906 576 0546 during 8 AM - 6 PM, Monday to Saturday.",
        "pidgin": "Call our Support people for support@trustnet.com or dial +234 906 576 0546 between 8 morning and 6 evening, Monday to Saturday.",
        "igbo": "Kpọtụrụ otu nkwado anyị na support@trustnet.com ma ọ bụ kpọọ +234 906 576 0546 n'oge 8 AM - 6 PM, Mọnde ruo Satọde.",
        "yoruba": "Pe ẹgbẹ Atilẹyin wa ni support@trustnet.com tabi pe +234 906 576 0546 laarin 8 owurọ si 6 osan, Monday si Saturday.",
        "hausa": "Tuntuɓi ƙungiyar tallafinmu a support@trustnet.com ko kiran +234 906 576 0546 daga 8 na safe zuwa 6 na yamma, Litinin zuwa Asabar."
    }
}

# Route for the root endpoint
@app.route("/")
def home():
    return "Backend is running!"

# Function to send WhatsApp message via API
def send_whatsapp_message(to, message):
    url = f"https://graph.facebook.com/v22.0/{WHATSAPP_PHONE_NUMBER_ID}/messages"
    headers = {
        "Authorization": f"Bearer {WHATSAPP_ACCESS_TOKEN}",
        "Content-Type": "application/json"
    }
    payload = {
        "messaging_product": "whatsapp",
        "to": to,
        "type": "text",
        "text": {"body": message}
    }
    requests.post(url, headers=headers, data=json.dumps(payload))

# Webhook verification endpoint (for WhatsApp integration)
@app.route("/webhook", methods=["GET"])
def verify_webhook():
    mode = request.args.get("hub.mode")
    token = request.args.get("hub.verify_token")
    challenge = request.args.get("hub.challenge")
    if mode == "subscribe" and token == VERIFY_TOKEN:
        return challenge, 200
    return "Forbidden", 403

# Webhook endpoint to handle incoming WhatsApp messages
@app.route("/webhook", methods=["POST"])
def webhook():
    data = request.get_json()
    print("Incoming message:", json.dumps(data, indent=2))

    if "entry" in data:
        for entry in data["entry"]:
            for change in entry.get("changes", []):
                value = change.get("value", {})
                messages = value.get("messages", [])
                for message in messages:
                    sender = message["from"]
                    text = message.get("text", {}).get("body", "").lower()

                    # Match keywords to responses
                    for keyword, replies in BOT_RESPONSES.items():
                        if keyword in text:
                            # Send replies in multiple languages
                            send_whatsapp_message(sender, replies["en"])
                            send_whatsapp_message(sender, replies["pidgin"])
                            send_whatsapp_message(sender, replies["igbo"])
                            send_whatsapp_message(sender, replies["yoruba"])
                            send_whatsapp_message(sender, replies["hausa"])
                            break
                    else:
                        send_whatsapp_message(sender, "I be Trustie. You fit ask me to Verify Documents, Report Scam, Know Your Rights, or Contact Support.")

    return "EVENT_RECEIVED", 200

# API Route for React to send data (document and language)
@app.route("/api/send", methods=["POST"])
def send_data():
    data = request.get_json()
    document = data.get("documentData")
    language = data.get("language", "en")
    
    # Here you can process the document or trigger WhatsApp messages
    recipient = data.get("phone")  # phone number from React
    if recipient:
        message = f"Your request is received! Language selected: {language}"
        send_whatsapp_message(recipient, message)

    return {"status": "success", "message": "Data received and processed"}, 200


# API Route for translation
@app.route("/api/translate", methods=["POST"])
async def translate_text():
    data = request.get_json()
    texts = data.get("texts")  # This should be a list
    target_lang = data.get("target_lang", "en")

    if not texts or not isinstance(texts, list):
        return {"success": False, "error": "No valid texts provided"}, 400

    try:
        translated_results = []
        for t in texts:
            translated = await translator.translate(t, dest=target_lang)
            translated_results.append({
                "original_text": t,
                "translated_text": translated.text,
                "source_lang": translated.src,
                "target_lang": target_lang
            })

        return {
            "success": True,
            "results": translated_results
        }
    except Exception as e:
        return {"success": False, "error": str(e)}, 500


# API Route for user registration
# ---------------- SIGN UP ----------------
@app.route("/signup", methods=["POST"])
def signup():
    data = request.get_json()

    # Required fields
    if not data.get("firstName") or not data.get("lastName") or not data.get("email") or not data.get("password"):
        return jsonify({"error": "All fields are required"}), 400

    # Validate names
    if not is_valid_name(data["firstName"]):
        return jsonify({"error": "Invalid first name"}), 400
    if not is_valid_name(data["lastName"]):
        return jsonify({"error": "Invalid last name"}), 400

    # Validate email
    if not is_valid_email(data["email"]):
        return jsonify({"error": "Invalid email format"}), 400

    # Check if email exists
    if User.query.filter_by(email=data["email"]).first():
        return jsonify({"error": "Email already exists"}), 400

    # Validate password strength
    if not is_strong_password(data["password"]):
        return jsonify({"error": "Password must be 8+ chars with upper, lower, digit, and special char"}), 400

    # Hash password
    hashed_password = bcrypt.generate_password_hash(data["password"]).decode("utf-8")

    # Save user
    new_user = User(
        first_name=data["firstName"],
        last_name=data["lastName"],
        email=data["email"],
        password=hashed_password,
    )
    db.session.add(new_user)
    db.session.commit()

    return jsonify({"message": "User registered successfully"}), 201

# ---------------- LOGIN ----------------
@app.route("/login", methods=["POST"])
def login():
    data = request.get_json()

    if not data.get("email") or not data.get("password"):
        return jsonify({"error": "Email and password required"}), 400

    user = User.query.filter_by(email=data["email"]).first()
    if not user or not bcrypt.check_password_hash(user.password, data["password"]):
        return jsonify({"error": "Invalid email or password"}), 401

    access_token = create_access_token(identity=user.id, expires_delta=timedelta(hours=1))
    return jsonify({"token": access_token, "user": { "id": user.id, "firstName": user.first_name, "lastName": user.last_name, "email": user.email}})

#---- forget password ----
@app.route("/forgot-password", methods=["POST"])
@limiter.limit("5 per hour")  # Max 5 requests per hour per IP
def forgot_password():
    data = request.get_json()
    email = data.get("email")

    if not email or not is_valid_email(email):
        return jsonify({"error": "Valid email required"}), 400

    user = User.query.filter_by(email=email).first()

    # 🛡 Do not reveal whether email exists
    if not user:
        return jsonify({"message": "If the email exists, a reset link has been sent."}), 200

    # Create reset token valid for 15 minutes
    reset_token = create_access_token(
        identity=user.id,
        expires_delta=timedelta(minutes=15),
        additional_claims={"reset": True}
    )

    reset_url = f"http://localhost:3000/reset-password/{reset_token}"

    msg = Message("Password Reset Request",
                  sender=app.config["MAIL_USERNAME"],
                  recipients=[email])
    msg.body = f"Click the link to reset your password: {reset_url}"
    mail.send(msg)

    return jsonify({"message": "If the email exists, a reset link has been sent."}), 200


# Function to verify survey plans
def verify_survey_plan(file_content, filename):
    """
    Simple survey plan verification
    Returns a score and detailed analysis
    """
    
    # Extract text from PDF
    text = ""
    try:
        if filename.lower().endswith('.pdf'):
            pdf_file = io.BytesIO(base64.b64decode(file_content))
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        
        elif filename.lower().endswith('.docx'):
            # Use python-docx for Word documents
            docx_file = io.BytesIO(base64.b64decode(file_content))
            doc = DocxDocument(docx_file)
            text = "\n".join([p.text for p in doc.paragraphs])

        else:
            # For other formats, assume it's already text
            text = file_content
            
        text = text.lower()  # Convert to lowercase for matching
        
    except Exception as e:
        return {
            "status": "error",
            "message": f"Could not process document: {str(e)}",
            "score": 0
        }
    
    # Define what we're looking for in a survey plan
    verification_checks = {
        "document_type": {
            "keywords": ["survey plan", "cadastral plan", "land survey", "survey document"],
            "found": False,
            "weight": 20
        },
        "surveyor_details": {
            "keywords": ["surveyor", "licensed", "registered", "number", "signature"],
            "found": False,
            "weight": 15
        },
        "property_details": {
            "keywords": ["parcel", "lot", "block", "plot", "plan no", "plan number"],
            "found": False,
            "weight": 15
        },
        "coordinates": {
            "keywords": ["coordinates", "latitude", "longitude", "bearing", "distance"],
            "found": False,
            "weight": 15
        },
        "scale": {
            "keywords": ["scale", "1:", "ratio"],
            "found": False,
            "weight": 10
        },
        "approval": {
            "keywords": ["approved", "government", "surveyor general", "stamped", "seal"],
            "found": False,
            "weight": 10
        },
        "date": {
            "keywords": ["date", "surveyed", "prepared", "year"],
            "found": False,
            "weight": 10
        },
        "boundaries": {
            "keywords": ["boundaries", "beacons", "pillars", "marks"],
            "found": False,
            "weight": 5
        }
    }
    
    # Check for each requirement
    total_score = 0
    max_score = sum(check["weight"] for check in verification_checks.values())
    found_elements = []
    missing_elements = []
    
    for element, check in verification_checks.items():
        found = any(keyword in text for keyword in check["keywords"])
        check["found"] = found
        
        if found:
            total_score += check["weight"]
            found_elements.append(element.replace("_", " ").title())
        else:
            missing_elements.append(element.replace("_", " ").title())
    
    # Calculate percentage score
    score = (total_score / max_score) * 100
    
    # Additional checks for red flags
    red_flags = []
    
    # Check for survey plan numbers
    plan_numbers = re.findall(r'plan\s*(no|number)?\s*[:]?\s*([A-Za-z0-9/-]+)', text, re.IGNORECASE)
    if not plan_numbers:
        red_flags.append("No survey plan number detected")
    
    # Check for suspicious patterns
    if "unofficial" in text or "copy" in text or "not for official use" in text:
        red_flags.append("Document contains unofficial markings")
    
    if len(text) < 300:  # Very short document
        red_flags.append("Document appears too short for a complete survey plan")
    
    # Check for required government stamps
    if "surveyor general" not in text and "government" not in text:
        red_flags.append("No government approval markings detected")
    
    # Determine status
    if score >= 85:
        status = "Verified"
        color = "green"
    elif score >= 65:
        status = "Caution"
        color = "orange"
    else:
        status = "Flagged"
        color = "red"
    
    return {
        "status": "success",
        "verification_result": {
            "score": round(score, 1),
            "status": status,
            "color": color,
            "found_elements": found_elements,
            "missing_elements": missing_elements,
            "red_flags": red_flags,
            "recommendations": generate_survey_recommendations(score, missing_elements, red_flags)
        }
    }

def generate_survey_recommendations(score, missing_elements, red_flags):
    """Generate recommendations based on survey verification results"""
    recommendations = []
    
    if score < 65:
        recommendations.append("Document requires professional surveyor review before proceeding")
    
    if "Document Type" in missing_elements:
        recommendations.append("Verify document is a legitimate survey plan")
    
    if "Surveyor Details" in missing_elements:
        recommendations.append("Confirm licensed surveyor details are present and valid")
    
    if "Property Details" in missing_elements:
        recommendations.append("Verify property identification details (lot, block, plan number)")
    
    if "Coordinates" in missing_elements:
        recommendations.append("Check that boundary coordinates are properly documented")
    
    if "Approval" in missing_elements:
        recommendations.append("Verify government approval stamps and signatures")
    
    if red_flags:
        recommendations.append("Address red flag issues before relying on this document")
    
    if score >= 85:
        recommendations.append("Document appears to meet basic survey plan requirements")
    
    return recommendations


# Endpoint to handle file upload and verification for survey plans
@app.route("/upload-and-verify-survey", methods=["POST"])
def upload_and_verify_survey():
    """
    Complete file upload and verification endpoint for survey plans
    Handles multipart file upload, processes immediately, and returns verification results
    """
    try:
        # Check if file is in request
        if 'file' not in request.files:
            return jsonify({"error": "No file uploaded"}), 400
        
        file = request.files['file']
        user_id = request.form.get('user_id')
        
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400
            
        if not user_id:
            return jsonify({"error": "User ID is required"}), 400
        
        # Check if user exists
        if not User.query.get(user_id):
            return jsonify({"error": "User not found"}), 404
        
        # Validate file type
        allowed_extensions = {'pdf', 'doc', 'docx', 'txt'}
        if not ('.' in file.filename and file.filename.rsplit('.', 1)[1].lower() in allowed_extensions):
            return jsonify({"error": "File type not supported. Use PDF, DOC, DOCX, or TXT"}), 400
        
        # Read file content
        file_content = file.read()
        
        # Generate file hash
        import hashlib
        file_hash = hashlib.sha256(file_content).hexdigest()
        
        
        # Convert file content to base64 for verification
        import base64
        file_content_b64 = base64.b64encode(file_content).decode('utf-8')
        
        # Run verification
        verification_result = verify_survey_plan(file_content_b64, file.filename)
        
        if verification_result["status"] == "error":
            return jsonify(verification_result), 400
        
        # Map verification status to your enum values
        verification_data = verification_result["verification_result"]
        status_mapping = {
            "Verified": "Verified",
            "Caution": "Caution", 
            "Flagged": "Flagged"
        }
        
        # Create document record
        new_document = Document(
            filename=file.filename,
            file_hash=file_hash,
            user_id=int(user_id),
            originality_status=status_mapping.get(verification_data["status"])
        )
        
        db.session.add(new_document)
        db.session.commit()
        
        # Return complete response with verification results
        return jsonify({
            "success": True,
            "message": "Survey plan uploaded and verified successfully",
            "document": {
                "id": new_document.id,
                "filename": new_document.filename,
                "file_hash": new_document.file_hash,
                "status": new_document.originality_status,
                "uploaded_at": new_document.uploaded_at.isoformat()
            },
            "verification": {
                "score": verification_data["score"],
                "status": verification_data["status"],
                "color": verification_data["color"],
                "found_elements": verification_data["found_elements"],
                "missing_elements": verification_data["missing_elements"],
                "red_flags": verification_data["red_flags"],
                "recommendations": verification_data["recommendations"],
                "summary": f"Survey plan scored {verification_data['score']}% and is classified as '{verification_data['status']}'"
            }
        }), 201
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": f"Upload and verification failed: {str(e)}"
        }), 500

# Get user documents with verification status
@app.route("/api/user-documents/<int:user_id>", methods=["GET"])
def get_user_documents(user_id):
    """Get all documents uploaded by a user with their verification status"""
    try:
        user = User.query.get(user_id)
        if not user:
            return jsonify({"error": "User not found"}), 404
        
        documents = Document.query.filter_by(user_id=user_id).order_by(Document.uploaded_at.desc()).all()
        
        documents_data = []
        for doc in documents:
            documents_data.append({
                "id": doc.id,
                "filename": doc.filename,
                "status": doc.originality_status,
                "uploaded_at": doc.uploaded_at.isoformat(),
                "file_hash": doc.file_hash[:16] + "..."  # Show partial hash for security
            })
        
        return jsonify({
            "success": True,
            "user": f"{user.first_name} {user.last_name}",
            "total_documents": len(documents_data),
            "documents": documents_data
        }), 200
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": f"Failed to fetch documents: {str(e)}"
        }), 500

# Get detailed verification results for a specific document
@app.route("/api/document-details/<int:document_id>", methods=["GET"])
def get_document_details(document_id):
    """Get detailed information about a specific document"""
    try:
        document = Document.query.get(document_id)
        if not document:
            return jsonify({"error": "Document not found"}), 404
        
        return jsonify({
            "success": True,
            "document": {
                "id": document.id,
                "filename": document.filename,
                "status": document.originality_status,
                "uploaded_at": document.uploaded_at.isoformat(),
                "uploaded_by": f"{document.user.first_name} {document.user.last_name}",
                "file_hash": document.file_hash
            }
        }), 200
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": f"Failed to fetch document details: {str(e)}"
        }), 500

# ============================================================================
# LAND TITLE VERIFICATION SYSTEM  
# ============================================================================

def verify_land_title_document(file_content, filename):
    """
    Verify Nigerian land title documents (C of O, Governor's Consent, Survey Plans)
    Returns a score and detailed analysis
    """
    
    # Extract text from document
    text = ""
    try:
        if filename.lower().endswith('.pdf'):
            pdf_file = io.BytesIO(base64.b64decode(file_content))
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        else:
            text = file_content
            
        text = text.lower()
        
    except Exception as e:
        return {
            "status": "error",
            "message": f"Could not process document: {str(e)}",
            "score": 0
        }
    
    # Define verification checks for land documents
    land_verification_checks = {
        "document_type": {
            "keywords": [
                "certificate of occupancy", "c of o", "c.of.o", "governor's consent", 
                "governors consent", "survey plan", "deed of assignment", "land title"
            ],
            "found": False,
            "weight": 25
        },
        "government_authority": {
            "keywords": [
                "lagos state government", "federal government", "ministry of lands", 
                "lands bureau", "governor", "surveyor general", "land registry", "Local Government"
            ],
            "found": False,
            "weight": 20
        },
        "property_description": {
            "keywords": [
                "plot", "block", "layout", "scheme", "estate", "coordinates", 
                "hectares", "square meters", "sq.m", "located at", "situate"
            ],
            "found": False,
            "weight": 15
        },
        "survey_details": {
            "keywords": [
                "survey", "surveyor", "beacon", "coordinates", "cadastral", 
                "triangulation", "survey plan no", "plan number", "surveyor's seal"
            ],
            "found": False,
            "weight": 15
        },
        "owner_details": {
            "keywords": [
                "grantee", "assignee", "owner", "proprietor", "holder", "occupier"
            ],
            "found": False,
            "weight": 10
        },
        "registration": {
            "keywords": [
                "registered", "registration", "file no", "file number", 
                "registry", "recorded", "gazetted", "excision"
            ],
            "found": False,
            "weight": 10
        },
        "signatures_seals": {
            "keywords": [
                "signature", "seal", "stamp", "commissioner", "permanent secretary", 
                "surveyor general", "official seal"
            ],
            "found": False,
            "weight": 5
        }
    }
    
    # Check for each requirement
    total_score = 0
    max_score = sum(check["weight"] for check in land_verification_checks.values())
    found_elements = []
    missing_elements = []
    
    for element, check in land_verification_checks.items():
        found = any(keyword in text for keyword in check["keywords"])
        check["found"] = found
        
        if found:
            total_score += check["weight"]
            found_elements.append(element.replace("_", " ").title())
        else:
            missing_elements.append(element.replace("_", " ").title())
    
    # Calculate percentage score
    score = (total_score / max_score) * 100
    
    # Red flags specific to land documents
    red_flags = []
    
    # Check for suspicious phrases
    warning_phrases = [
        "excision in progress", "pending excision", "family land", "community land",
        "customary right", "under litigation", "court case", "disputed", 
        "temporary", "provisional", "interim"
    ]
    
    for phrase in warning_phrases:
        if phrase in text:
            red_flags.append(f"Contains warning phrase: '{phrase}'")
    
    # Check for missing critical elements
    if "certificate of occupancy" not in text and "c of o" not in text and "governor's consent" not in text:
        red_flags.append("Does not appear to be a recognized land title document")
    
    if len(text) < 300:
        red_flags.append("Document appears too short for a complete land title")
    
    # Check for proper formatting indicators
    if "file no" not in text and "registration" not in text:
        red_flags.append("Missing registration or file number")
    
    # Determine document status
    if score >= 75:
        status = "Verified"
        color = "green"
    elif score >= 50:
        status = "Caution"
        color = "orange"
    else:
        status = "Flagged"
        color = "red"
    
    return {
        "status": "success",
        "verification_result": {
            "document_type": "land_title",
            "score": round(score, 1),
            "status": status,
            "color": color,
            "found_elements": found_elements,
            "missing_elements": missing_elements,
            "red_flags": red_flags,
            "recommendations": generate_land_recommendations(score, missing_elements, red_flags, text)
        }
    }

def generate_land_recommendations(score, missing_elements, red_flags, text):
    """Generate recommendations for land title documents"""
    recommendations = []
    
    if score < 50:
        recommendations.append("Document requires thorough verification with state Lands Registry")
    
    if "Document Type" in missing_elements:
        recommendations.append("Verify this is a legitimate land title document")
    
    if "Government Authority" in missing_elements:
        recommendations.append("Confirm document is issued by proper government authority")
    
    if "Survey Details" in missing_elements:
        recommendations.append("Ensure survey plan details are included and verified")
    
    if "Registration" in missing_elements:
        recommendations.append("Verify document is properly registered with Lands Registry")
    
    # Nigerian-specific land verification advice
    recommendations.append("Always verify at the state Lands Registry before purchase")
    recommendations.append("Conduct a search at the Corporate Affairs Commission if dealing with corporate entities")
    
    if "excision" in text.lower():
        recommendations.append("For excision lands, ensure excision process is complete and gazetted")
    
    if "survey plan" in text.lower():
        recommendations.append("Verify survey plan with Surveyor General's office")
    
    if red_flags:
        recommendations.append("Address all red flag issues before proceeding with transaction")
    
    if score >= 75:
        recommendations.append("Document appears authentic but still requires registry verification")
    else:
        recommendations.append("Consider engaging a qualified estate surveyor for verification")
    
    return recommendations

# ============================================================================
# SEPARATE API ROUTES FOR EACH DOCUMENT TYPE
# ============================================================================


# Route specifically for land title upload and verification
@app.route("/upload-land-title", methods=["POST"])
def upload_land_title():
    """Upload and verify land title documents (C of O, Governor's Consent, etc.)"""
    try:
        # Check if file is in request
        if 'file' not in request.files:
            return jsonify({"error": "No file uploaded"}), 400
        
        file = request.files['file']
        user_id = request.form.get('user_id')
        
        if file.filename == '' or not user_id:
            return jsonify({"error": "File and user ID are required"}), 400
        
        # Check if user exists
        if not User.query.get(user_id):
            return jsonify({"error": "User not found"}), 404
        
        # Validate file type
        allowed_extensions = {'pdf', 'doc', 'docx', 'txt'}
        if not ('.' in file.filename and file.filename.rsplit('.', 1)[1].lower() in allowed_extensions):
            return jsonify({"error": "File type not supported. Use PDF, DOC, DOCX, or TXT"}), 400
        
        # Read and process file
        file_content = file.read()
        file_hash = hashlib.sha256(file_content).hexdigest()
        
        
        # Convert to base64 for verification
        file_content_b64 = base64.b64encode(file_content).decode('utf-8')
        
        # Run land title verification
        verification_result = verify_land_title_document(file_content_b64, file.filename)
        
        if verification_result["status"] == "error":
            return jsonify(verification_result), 400
        
        # Map status to your enum
        verification_data = verification_result["verification_result"]
        status_mapping = {
            "Likely Authentic": "Verified",
            "Needs Verification": "Caution", 
            "Suspicious": "Flagged"
        }
        
        # Save to database
        new_document = Document(
            filename=file.filename,
            file_hash=file_hash,
            user_id=int(user_id),
            originality_status=status_mapping.get(verification_data["status"])
        )
        
        db.session.add(new_document)
        db.session.commit()
        
        return jsonify({
            "success": True,
            "message": "Land title document uploaded and verified successfully",
            "document_type": "Land Title Document",
            "document": {
                "id": new_document.id,
                "filename": new_document.filename,
                "status": new_document.originality_status,
                "uploaded_at": new_document.uploaded_at.isoformat()
            },
            "verification": verification_data
        }), 201
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": f"Upload failed: {str(e)}"
        }), 500

# route for chat functionality
@app.route('/chat', methods=['POST'])
def chat():
    try:
        data = request.get_json()
        user_message = data.get('message', '')
        logging.info(f"Received chat message: {user_message}")

        kb_answer = search_knowledge_base(user_message)
        if kb_answer:
            logging.info(f"Knowledge base answer found: {kb_answer}")
            return jsonify({'response': kb_answer})
        

        response = get_gemini_response(user_message)
        return jsonify({'response': response})
    
    except Exception as e:
        logging.error(f"Error in chat(): {e}")
        logging.error(traceback.format_exc()) 
        return jsonify({'response': "An error occurred while processing your message."})

            
if __name__ == "__main__":
    app.run(port=5000, debug=True)